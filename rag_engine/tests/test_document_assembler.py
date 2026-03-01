from __future__ import annotations
import sys
import os

sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..'))

import pytest
from docx import Document as DocxDocument
from docx.shared import Pt
from document_assembler import assemble_docx, KRDS_FONT, _BLUE_900


def test_assemble_basic_docx(tmp_path):
    sections = [
        ("제안 개요", "## 제안 개요\n\n본 사업은 XX기관의 정보시스템 구축을 위한 제안입니다."),
        ("기술적 접근방안", "## 기술적 접근방안\n\n### 시스템 아키텍처\n\n3-tier 구조로 구성합니다.\n\n- 웹서버: Nginx\n- WAS: Spring Boot\n- DB: PostgreSQL"),
    ]
    out_path = str(tmp_path / "proposal.docx")
    result = assemble_docx(
        title="XX기관 정보시스템 구축 제안서",
        sections=sections,
        output_path=out_path,
    )
    assert os.path.exists(result)
    assert result.endswith(".docx")
    assert os.path.getsize(result) > 1000


def test_assemble_empty_sections(tmp_path):
    out_path = str(tmp_path / "empty.docx")
    result = assemble_docx(
        title="빈 제안서",
        sections=[],
        output_path=out_path,
    )
    assert os.path.exists(result)


def test_assemble_krds_styles_applied(tmp_path):
    """KRDS 디자인 토큰이 DOCX에 적용되는지 확인."""
    sections = [
        ("사업 이해", "## 사업 이해\n\n본 사업의 핵심 목적을 설명합니다."),
        ("기술 방안", "## 기술 방안\n\n- React\n- FastAPI\n- PostgreSQL"),
    ]
    out_path = str(tmp_path / "krds.docx")
    assemble_docx(
        title="스마트시티 구축 제안서",
        sections=sections,
        output_path=out_path,
        company_name="(주)MS솔루션즈",
    )

    doc = DocxDocument(out_path)

    # Normal style should use KRDS font
    normal = doc.styles["Normal"]
    assert normal.font.name == KRDS_FONT
    assert normal.font.size == Pt(11)

    # Heading 1 should use Blue 900
    h1 = doc.styles["Heading 1"]
    assert h1.font.color.rgb == _BLUE_900
    assert h1.font.size == Pt(18)
    assert h1.font.bold is True

    # Left margin 3.0cm for binding
    from docx.shared import Cm
    section = doc.sections[0]
    assert section.left_margin is not None
    assert section.left_margin >= Cm(3.0)

    # Line spacing 1.25
    assert normal.paragraph_format.line_spacing == 1.25

    # Footer should contain company name
    footer_text = "".join(
        p.text for p in doc.sections[0].footer.paragraphs
    )
    assert "MS솔루션즈" in footer_text

    # Cover page — find title text
    all_text = [p.text for p in doc.paragraphs]
    assert any("스마트시티" in t for t in all_text)
    assert any("CONFIDENTIAL" in t for t in all_text)

    # Cover page should have date (YYYY년 MM월)
    import re
    assert any(re.search(r"\d{4}년 \d{2}월", t) for t in all_text)

    # TOC page — numbered entries with dot leaders
    assert any("01" in t and "사업 이해" in t for t in all_text)


def test_assemble_preserves_inline_formatting(tmp_path):
    """Bold/italic markdown formatting is preserved in DOCX runs."""
    sections = [
        ("테스트", "## 테스트\n\n이것은 **굵은** 텍스트와 *기울임* 텍스트입니다."),
    ]
    out_path = str(tmp_path / "inline.docx")
    assemble_docx(
        title="인라인 포맷 테스트",
        sections=sections,
        output_path=out_path,
    )

    doc = DocxDocument(out_path)
    all_runs = []
    for p in doc.paragraphs:
        for run in p.runs:
            all_runs.append((run.text, run.bold, run.italic))

    # Should have a bold run with "굵은"
    bold_runs = [r for r in all_runs if r[0] == "굵은" and r[1] is True]
    assert len(bold_runs) >= 1, f"Expected bold '굵은' run, got: {all_runs}"

    # Should have an italic run with "기울임"
    italic_runs = [r for r in all_runs if r[0] == "기울임" and r[2] is True]
    assert len(italic_runs) >= 1, f"Expected italic '기울임' run, got: {all_runs}"
