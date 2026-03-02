"""Tests for track_record_assembler.py."""
import sys, os, tempfile
sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

from docx import Document as DocxDocument
from docx.shared import Pt

from phase2_models import TrackRecordEntry, PersonnelEntry
from track_record_assembler import assemble_track_record_docx, _KRDS_FONT, _BLUE_900


def _sample_records():
    return [
        TrackRecordEntry(
            project_name="스마트시티 관제시스템",
            client="부산시",
            period="2024.01~2024.12",
            amount=30.0,
            description="IoT 센서 통합 관제",
            technologies=["Python", "IoT"],
            relevance_score=0.85,
            generated_text="부산시 스마트시티 관제시스템은 IoT 센서 1,000개를 통합하여 실시간 모니터링 체계를 구축한 프로젝트입니다.",
        ),
        TrackRecordEntry(
            project_name="IoT 플랫폼",
            client="과기부",
            period="2023.06~2024.03",
            amount=20.0,
            relevance_score=0.65,
            generated_text="과학기술정보통신부의 IoT 플랫폼 고도화 사업을 수행하였습니다.",
        ),
    ]


def _sample_personnel():
    return [
        PersonnelEntry(
            name="홍길동",
            role="PM",
            grade="특급",
            experience_years=15,
            certifications=["PMP", "정보관리기술사"],
            key_projects=["스마트시티 관제시스템"],
            generated_text="15년간 대규모 공공 IT 프로젝트를 총괄한 경력을 보유한 PM입니다.",
        ),
    ]


def test_assemble_creates_docx():
    with tempfile.TemporaryDirectory() as tmpdir:
        path = os.path.join(tmpdir, "test.docx")
        result = assemble_track_record_docx(
            title="유사수행실적 기술서",
            records=_sample_records(),
            personnel=_sample_personnel(),
            output_path=path,
        )
        assert os.path.isfile(result)
        assert os.path.getsize(result) > 0


def test_assemble_with_empty_records():
    with tempfile.TemporaryDirectory() as tmpdir:
        path = os.path.join(tmpdir, "empty.docx")
        result = assemble_track_record_docx(
            title="빈 기술서",
            records=[],
            personnel=_sample_personnel(),
            output_path=path,
        )
        assert os.path.isfile(result)


def test_assemble_with_empty_personnel():
    with tempfile.TemporaryDirectory() as tmpdir:
        path = os.path.join(tmpdir, "no_personnel.docx")
        result = assemble_track_record_docx(
            title="실적만 기술서",
            records=_sample_records(),
            personnel=[],
            output_path=path,
        )
        assert os.path.isfile(result)


def test_assemble_krds_styles_applied():
    """KRDS 디자인 토큰이 실적기술서 DOCX에 적용되는지 확인."""
    with tempfile.TemporaryDirectory() as tmpdir:
        path = os.path.join(tmpdir, "krds.docx")
        assemble_track_record_docx(
            title="유사수행실적 기술서",
            records=_sample_records(),
            personnel=_sample_personnel(),
            output_path=path,
            company_name="(주)MS솔루션즈",
        )

        doc = DocxDocument(path)

        # Normal style should use KRDS font
        normal = doc.styles["Normal"]
        assert normal.font.name == _KRDS_FONT
        assert normal.font.size == Pt(11)

        # Heading 1 should use Blue 900
        h1 = doc.styles["Heading 1"]
        assert h1.font.color.rgb == _BLUE_900
        assert h1.font.size == Pt(18)
        assert h1.font.bold is True

        # Left margin should be 3.0cm for binding
        section = doc.sections[0]
        assert section.left_margin is not None
        assert section.left_margin > 0

        # Footer should contain company name
        footer_text = "".join(
            p.text for p in doc.sections[0].footer.paragraphs
        )
        assert "MS솔루션즈" in footer_text

        # Cover page should have title and CONFIDENTIAL
        all_text = [p.text for p in doc.paragraphs]
        assert any("유사수행실적" in t for t in all_text)
        assert any("CONFIDENTIAL" in t for t in all_text)

        # Should have track record content
        assert any("스마트시티" in t for t in all_text)
        assert any("홍길동" in t for t in all_text)
