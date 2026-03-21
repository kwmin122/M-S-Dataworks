"""WBS Generator — XLSX + 간트차트 + DOCX 생성.

openpyxl로 WBS/인력배치표/산출물 3시트 XLSX,
matplotlib로 간트차트 PNG,
python-docx로 수행계획서 DOCX 생성.

KRDS 디자인 시스템 토큰 적용:
- Blue 900 (#003764) 헤더/제목
- Pretendard 폰트 (없으면 맑은 고딕 fallback)
"""
from __future__ import annotations

import logging
import os
from typing import Optional

from phase2_models import WbsTask, PersonnelAllocation

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# KRDS Design Tokens
# ---------------------------------------------------------------------------

_KRDS_BLUE_900 = "003764"
_KRDS_BLUE_800 = "004A8F"
_KRDS_BLUE_600 = "0070E0"
_KRDS_BLUE_100 = "E6F0FF"
_KRDS_BLUE_50 = "F0F6FF"
_KRDS_FONT = "Pretendard"
_FALLBACK_FONT = "맑은 고딕"


# ---------------------------------------------------------------------------
# XLSX 생성
# ---------------------------------------------------------------------------

def generate_wbs_xlsx(
    tasks: list[WbsTask],
    personnel: list[PersonnelAllocation],
    title: str,
    total_months: int,
    output_path: str,
) -> str:
    """WBS XLSX 생성 (3시트: WBS / 인력배치표 / 산출물)."""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    wb = Workbook()

    # Style definitions (KRDS)
    header_fill = PatternFill(start_color=_KRDS_BLUE_900, end_color=_KRDS_BLUE_900, fill_type="solid")
    header_font = Font(name=_KRDS_FONT, size=10, bold=True, color="FFFFFF")
    cell_font = Font(name=_KRDS_FONT, size=9)
    center_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
    thin_border = Border(
        left=Side(style="thin"),
        right=Side(style="thin"),
        top=Side(style="thin"),
        bottom=Side(style="thin"),
    )
    phase_fill = PatternFill(start_color=_KRDS_BLUE_100, end_color=_KRDS_BLUE_100, fill_type="solid")

    # --- Sheet 1: WBS ---
    ws_wbs = wb.active
    ws_wbs.title = "WBS"

    # Header
    wbs_headers = ["No.", "단계", "작업명", "담당", "시작월", "종료월", "기간(월)", "M/M"]
    # Add month columns
    for m in range(1, total_months + 1):
        wbs_headers.append(f"{m}월")

    for col_idx, header in enumerate(wbs_headers, 1):
        cell = ws_wbs.cell(row=1, column=col_idx, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = center_align
        cell.border = thin_border

    # Data rows
    current_phase = ""
    for idx, task in enumerate(tasks, 1):
        row = idx + 1
        ws_wbs.cell(row=row, column=1, value=idx).font = cell_font
        ws_wbs.cell(row=row, column=2, value=task.phase).font = cell_font
        ws_wbs.cell(row=row, column=3, value=task.task_name).font = cell_font
        ws_wbs.cell(row=row, column=4, value=task.responsible_role).font = cell_font
        ws_wbs.cell(row=row, column=5, value=task.start_month).font = cell_font
        ws_wbs.cell(row=row, column=6, value=task.end_month()).font = cell_font
        ws_wbs.cell(row=row, column=7, value=task.duration_months).font = cell_font
        ws_wbs.cell(row=row, column=8, value=task.man_months).font = cell_font

        for col in range(1, len(wbs_headers) + 1):
            ws_wbs.cell(row=row, column=col).border = thin_border
            ws_wbs.cell(row=row, column=col).alignment = center_align

        # Phase row shading
        if task.phase != current_phase:
            current_phase = task.phase
            for col in range(1, 9):
                ws_wbs.cell(row=row, column=col).fill = phase_fill

        # Gantt bars in month columns (KRDS Blue 600)
        gantt_fill = PatternFill(start_color=_KRDS_BLUE_600, end_color=_KRDS_BLUE_600, fill_type="solid")
        for m in range(task.start_month, min(task.end_month() + 1, total_months + 1)):
            col = 8 + m  # month columns start at column 9
            ws_wbs.cell(row=row, column=col).fill = gantt_fill

    # Column widths
    ws_wbs.column_dimensions["A"].width = 5
    ws_wbs.column_dimensions["B"].width = 12
    ws_wbs.column_dimensions["C"].width = 25
    ws_wbs.column_dimensions["D"].width = 10
    ws_wbs.column_dimensions["E"].width = 8
    ws_wbs.column_dimensions["F"].width = 8
    ws_wbs.column_dimensions["G"].width = 8
    ws_wbs.column_dimensions["H"].width = 8
    # Month columns
    from openpyxl.utils import get_column_letter
    for m in range(1, total_months + 1):
        col_letter = get_column_letter(8 + m)
        ws_wbs.column_dimensions[col_letter].width = 6

    # --- Sheet 2: 인력배치표 ---
    ws_staff = wb.create_sheet("인력배치표")
    staff_headers = ["No.", "역할", "등급", "합계(M/M)"]
    for m in range(1, total_months + 1):
        staff_headers.append(f"{m}월")

    for col_idx, header in enumerate(staff_headers, 1):
        cell = ws_staff.cell(row=1, column=col_idx, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = center_align
        cell.border = thin_border

    for idx, alloc in enumerate(personnel, 1):
        row = idx + 1
        ws_staff.cell(row=row, column=1, value=idx).font = cell_font
        ws_staff.cell(row=row, column=2, value=alloc.role).font = cell_font
        ws_staff.cell(row=row, column=3, value=alloc.grade).font = cell_font
        ws_staff.cell(row=row, column=4, value=alloc.total_man_months).font = cell_font

        for m_idx, mm in enumerate(alloc.monthly_allocation):
            col = 5 + m_idx
            cell = ws_staff.cell(row=row, column=col, value=round(mm, 1) if mm > 0 else "")
            cell.font = cell_font
            cell.alignment = center_align
            cell.border = thin_border

        for col in range(1, len(staff_headers) + 1):
            ws_staff.cell(row=row, column=col).border = thin_border
            ws_staff.cell(row=row, column=col).alignment = center_align

    # --- Sheet 3: 산출물 ---
    ws_del = wb.create_sheet("산출물")
    del_headers = ["No.", "단계", "산출물명", "산출 시기"]

    for col_idx, header in enumerate(del_headers, 1):
        cell = ws_del.cell(row=1, column=col_idx, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = center_align
        cell.border = thin_border

    del_idx = 1
    for task in tasks:
        for deliverable in task.deliverables:
            row = del_idx + 1
            ws_del.cell(row=row, column=1, value=del_idx).font = cell_font
            ws_del.cell(row=row, column=2, value=task.phase).font = cell_font
            ws_del.cell(row=row, column=3, value=deliverable).font = cell_font
            ws_del.cell(row=row, column=4, value=f"{task.end_month()}월").font = cell_font
            for col in range(1, 5):
                ws_del.cell(row=row, column=col).border = thin_border
                ws_del.cell(row=row, column=col).alignment = center_align
            del_idx += 1

    ws_del.column_dimensions["C"].width = 30

    wb.save(output_path)
    return output_path


# ---------------------------------------------------------------------------
# 간트차트 PNG 생성
# ---------------------------------------------------------------------------

def _find_korean_font() -> Optional[str]:
    """Find Korean font with robust fallback for all platforms."""
    candidates = [
        # macOS
        "/System/Library/Fonts/AppleSDGothicNeo.ttc",
        "/Library/Fonts/NanumGothic.ttf",
        # Linux (Docker/Railway)
        "/usr/share/fonts/truetype/nanum/NanumGothic.ttf",
        "/usr/share/fonts/nanum/NanumGothic.ttf",
        "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/google-noto-cjk/NotoSansCJK-Regular.ttc",
        # Windows
        "C:/Windows/Fonts/malgun.ttf",
        "C:/Windows/Fonts/NanumGothic.ttf",
    ]
    for path in candidates:
        if os.path.isfile(path):
            return path

    # Last resort: search common font directories
    import glob
    for pattern in [
        "/usr/share/fonts/**/Nanum*.ttf",
        "/usr/share/fonts/**/NotoSans*CJK*.ttc",
        "/usr/share/fonts/**/NotoSans*CJK*.otf",
    ]:
        matches = glob.glob(pattern, recursive=True)
        if matches:
            return matches[0]

    return None


def generate_gantt_chart(
    tasks: list[WbsTask],
    total_months: int,
    output_path: str,
) -> str:
    """matplotlib 간트차트 이미지 생성."""
    import matplotlib
    matplotlib.use("Agg")  # non-interactive backend
    import matplotlib.pyplot as plt
    from matplotlib.patches import FancyBboxPatch
    import matplotlib.font_manager as fm

    # Korean font setup
    ko_font_path = _find_korean_font()
    if ko_font_path:
        font_prop = fm.FontProperties(fname=ko_font_path, size=8)
        title_font = fm.FontProperties(fname=ko_font_path, size=12, weight="bold")
    else:
        logger.warning("Korean font not found — Gantt labels may render incorrectly")
        font_prop = fm.FontProperties(size=8)
        title_font = fm.FontProperties(size=12, weight="bold")

    fig_height = max(6, len(tasks) * 0.4 + 2)
    fig, ax = plt.subplots(figsize=(14, fig_height))

    # Color palette per phase (KRDS)
    phase_colors = {}
    color_palette = [
        "#003764", "#004A8F", "#005DB8", "#0070E0",
        "#1A85FF", "#0D8050", "#E07B54", "#0D6E6E",
    ]
    phase_idx = 0
    for task in tasks:
        if task.phase not in phase_colors:
            phase_colors[task.phase] = color_palette[phase_idx % len(color_palette)]
            phase_idx += 1

    # Draw bars
    y_positions = list(range(len(tasks) - 1, -1, -1))
    for i, task in enumerate(tasks):
        y = y_positions[i]
        color = phase_colors[task.phase]
        bar = FancyBboxPatch(
            (task.start_month - 0.5, y - 0.3),
            task.duration_months,
            0.6,
            boxstyle="round,pad=0.05",
            facecolor=color,
            edgecolor="white",
            linewidth=0.5,
            alpha=0.85,
        )
        ax.add_patch(bar)

    # Y-axis labels
    labels = [f"{t.phase} | {t.task_name}" for t in tasks]
    ax.set_yticks(y_positions)
    ax.set_yticklabels(labels, fontproperties=font_prop)

    # X-axis
    ax.set_xlim(0.3, total_months + 0.7)
    ax.set_xticks(range(1, total_months + 1))
    ax.set_xticklabels([f"{m}월" for m in range(1, total_months + 1)], fontproperties=font_prop)

    ax.set_xlabel("사업 기간", fontproperties=font_prop)
    ax.set_title("프로젝트 일정 (간트차트)", fontproperties=title_font, pad=15)

    # Grid
    ax.set_axisbelow(True)
    ax.xaxis.grid(True, linestyle="--", alpha=0.3)
    ax.set_ylim(-0.5, len(tasks) - 0.5)

    plt.tight_layout()
    fig.savefig(output_path, dpi=150, bbox_inches="tight", facecolor="white")
    plt.close(fig)

    return output_path


# ---------------------------------------------------------------------------
# 수행계획서 DOCX 생성
# ---------------------------------------------------------------------------

def _setup_wbs_docx_styles(doc) -> None:
    """KRDS 디자인 토큰을 WBS DOCX에 적용."""
    from docx.shared import Pt, Cm, RGBColor
    from docx.oxml.ns import qn

    blue_900 = RGBColor(0x00, 0x37, 0x64)
    blue_800 = RGBColor(0x00, 0x4A, 0x8F)
    gray_900 = RGBColor(0x1A, 0x1A, 0x1A)
    gray_700 = RGBColor(0x44, 0x44, 0x44)

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)   # 바인딩 여백
        section.right_margin = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = _KRDS_FONT
    style.font.size = Pt(11)
    style.font.color.rgb = gray_700
    pf = style.paragraph_format
    pf.space_before = Pt(3)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.25
    rpr = style.element.get_or_add_rPr()
    ea = rpr.find(qn("w:rFonts"))
    if ea is None:
        ea = style.element.makeelement(qn("w:rFonts"), {})
        rpr.insert(0, ea)
    ea.set(qn("w:eastAsia"), _KRDS_FONT)

    for level in range(1, 4):
        hstyle = doc.styles[f"Heading {level}"]
        sizes = {1: 18, 2: 14, 3: 12}
        colors = {1: blue_900, 2: blue_800, 3: gray_900}
        hstyle.font.name = _KRDS_FONT
        hstyle.font.size = Pt(sizes[level])
        hstyle.font.bold = True
        hstyle.font.color.rgb = colors[level]
        hrpr = hstyle.element.get_or_add_rPr()
        hea = hrpr.find(qn("w:rFonts"))
        if hea is None:
            hea = hstyle.element.makeelement(qn("w:rFonts"), {})
            hrpr.insert(0, hea)
        hea.set(qn("w:eastAsia"), _KRDS_FONT)


def _style_wbs_table_header(row, col_count: int) -> None:
    """KRDS Blue 900 배경 + 흰색 텍스트로 표 헤더 스타일링."""
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml

    for i in range(col_count):
        cell = row.cells[i]
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{_KRDS_BLUE_900}"/>')
        cell._tc.get_or_add_tcPr().append(shading)
        # Re-create text with explicit runs for reliable styling
        existing_text = cell.text
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(existing_text)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.name = _KRDS_FONT
        rpr = run._element.get_or_add_rPr()
        ea = rpr.find(qn("w:rFonts"))
        if ea is None:
            ea = run._element.makeelement(qn("w:rFonts"), {})
            rpr.insert(0, ea)
        ea.set(qn("w:eastAsia"), _KRDS_FONT)


def generate_wbs_docx(
    tasks: list[WbsTask],
    personnel: list[PersonnelAllocation],
    title: str,
    total_months: int,
    methodology_name: str,
    output_path: str,
    gantt_path: Optional[str] = None,
    company_name: str = "",
) -> str:
    """수행계획서 DOCX 생성 (WBS표 + 인력표 + 간트차트 삽입)."""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn

    doc = Document()
    doc.core_properties.author = "Kira Bot"
    _setup_wbs_docx_styles(doc)

    blue_900 = RGBColor(0x00, 0x37, 0x64)
    blue_600 = RGBColor(0x00, 0x70, 0xE0)
    gray_500 = RGBColor(0x88, 0x88, 0x88)
    gray_900 = RGBColor(0x1A, 0x1A, 0x1A)

    # Cover page
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.size = Pt(24)
    run.bold = True
    run.font.color.rgb = blue_900
    run.font.name = _KRDS_FONT

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("━" * 20)
    run2.font.color.rgb = blue_600
    run2.font.size = Pt(12)

    if company_name:
        for _ in range(3):
            doc.add_paragraph()
        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run3 = p3.add_run(company_name)
        run3.font.size = Pt(16)
        run3.bold = True
        run3.font.color.rgb = gray_900
        run3.font.name = _KRDS_FONT

    # Date (KST)
    from datetime import datetime, timezone, timedelta
    kst = timezone(timedelta(hours=9))
    now = datetime.now(kst)
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_date = p_date.add_run(now.strftime("%Y년 %m월"))
    run_date.font.size = Pt(11)
    run_date.font.color.rgb = gray_500
    run_date.font.name = _KRDS_FONT

    doc.add_page_break()

    # 1. 수행방법론
    doc.add_heading("제1장 수행방법론", level=1)
    doc.add_heading(f"1.1 {methodology_name} 방법론", level=2)
    doc.add_paragraph(
        f"본 사업은 {methodology_name} 방법론을 적용하여 총 {total_months}개월간 수행합니다. "
        f"전체 {len(tasks)}개 세부 태스크를 통해 체계적으로 프로젝트를 관리합니다."
    )

    # 2. WBS 요약표
    doc.add_heading("제2장 WBS (작업분류체계)", level=1)

    # Phase summary
    phase_tasks: dict[str, list[WbsTask]] = {}
    for task in tasks:
        phase_tasks.setdefault(task.phase, []).append(task)

    for phase_name, phase_task_list in phase_tasks.items():
        doc.add_heading(phase_name, level=2)
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        headers = ["작업명", "담당", "기간", "M/M", "산출물"]
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
        _style_wbs_table_header(table.rows[0], 5)

        for task in phase_task_list:
            row = table.add_row()
            row.cells[0].text = task.task_name
            row.cells[1].text = task.responsible_role
            row.cells[2].text = f"{task.start_month}~{task.end_month()}월"
            row.cells[3].text = str(task.man_months)
            row.cells[4].text = ", ".join(task.deliverables) if task.deliverables else "-"

        doc.add_paragraph()

    # 3. 간트차트 삽입
    if gantt_path and os.path.isfile(gantt_path):
        doc.add_heading("제3장 프로젝트 일정", level=1)
        doc.add_picture(gantt_path, width=Cm(16))
        doc.add_paragraph()

    # 4. 투입인력
    doc.add_heading("제4장 투입인력 계획", level=1)
    if personnel:
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(["역할", "등급", "합계(M/M)", "비고"]):
            table.rows[0].cells[i].text = h
        _style_wbs_table_header(table.rows[0], 4)

        total_mm = 0.0
        for alloc in personnel:
            row = table.add_row()
            row.cells[0].text = alloc.role
            row.cells[1].text = alloc.grade
            row.cells[2].text = str(alloc.total_man_months)
            row.cells[3].text = ""
            total_mm += alloc.total_man_months

        # Total row
        row = table.add_row()
        row.cells[0].text = "합계"
        row.cells[2].text = str(round(total_mm, 1))

    # Footer
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if company_name:
            fr = fp.add_run(f"{company_name}  |  CONFIDENTIAL")
        else:
            fr = fp.add_run("CONFIDENTIAL")
        fr.font.size = Pt(8)
        fr.font.color.rgb = gray_500
        fr.font.name = _KRDS_FONT

    doc.save(output_path)
    return output_path
