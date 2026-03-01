"""Document Assembler — DOCX output from proposal sections.

Takes list of section texts (markdown) and assembles a .docx file
with proper formatting using python-docx. Markdown is parsed via
mistune AST renderer for reliable heading/list extraction.

KRDS 디자인 시스템 기반 포맷:
- Pretendard 폰트 (없으면 맑은 고딕 fallback)
- Blue 900 (#003764) 기반 제목/헤딩
- Gray 700 (#444444) 본문
- 표지/목차/섹션 구분 스타일링
"""
from __future__ import annotations

import re
from datetime import datetime, timezone, timedelta
from typing import Any

import mistune
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

# mistune 3.x AST renderer — produces a token tree we walk to emit DOCX elements.
_md_parser = mistune.create_markdown(renderer="ast")

# ---------------------------------------------------------------------------
# KRDS Design Tokens for DOCX
# ---------------------------------------------------------------------------

KRDS_FONT = "Pretendard"
_FALLBACK_FONT = "맑은 고딕"

_BLUE_900 = RGBColor(0x00, 0x37, 0x64)
_BLUE_800 = RGBColor(0x00, 0x4A, 0x8F)
_BLUE_600 = RGBColor(0x00, 0x70, 0xE0)
_BLUE_100 = RGBColor(0xE6, 0xF0, 0xFF)
_BLUE_50 = RGBColor(0xF0, 0xF6, 0xFF)
_GRAY_900 = RGBColor(0x1A, 0x1A, 0x1A)
_GRAY_700 = RGBColor(0x44, 0x44, 0x44)
_GRAY_500 = RGBColor(0x88, 0x88, 0x88)
_GRAY_200 = RGBColor(0xE5, 0xE5, 0xE5)
_WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def _set_font(run, size: float, bold: bool = False, color: RGBColor = _GRAY_700):
    """Apply KRDS font settings to a run."""
    run.font.name = KRDS_FONT
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color
    # CJK font fallback
    rpr = run._element.get_or_add_rPr()
    ea = rpr.find(qn("w:rFonts"))
    if ea is None:
        ea = run._element.makeelement(qn("w:rFonts"), {})
        rpr.insert(0, ea)
    ea.set(qn("w:eastAsia"), KRDS_FONT)


def _setup_styles(doc: Document) -> None:
    """Configure document-level styles for KRDS compliance."""
    # Page margins (left 3.0cm for binding)
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)
        section.header_distance = Cm(1.2)
        section.footer_distance = Cm(1.0)

    # Default paragraph style
    style = doc.styles["Normal"]
    style.font.name = KRDS_FONT
    style.font.size = Pt(11)
    style.font.color.rgb = _GRAY_700
    pf = style.paragraph_format
    pf.space_before = Pt(3)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.25
    rpr = style.element.get_or_add_rPr()
    ea = rpr.find(qn("w:rFonts"))
    if ea is None:
        ea = style.element.makeelement(qn("w:rFonts"), {})
        rpr.insert(0, ea)
    ea.set(qn("w:eastAsia"), KRDS_FONT)

    # Heading styles
    for level in range(1, 4):
        hstyle = doc.styles[f"Heading {level}"]
        sizes = {1: 18, 2: 14, 3: 12}
        colors = {1: _BLUE_900, 2: _BLUE_800, 3: _GRAY_900}
        hstyle.font.name = KRDS_FONT
        hstyle.font.size = Pt(sizes[level])
        hstyle.font.bold = True
        hstyle.font.color.rgb = colors[level]
        hpf = hstyle.paragraph_format
        hpf.space_before = Pt(18 if level == 1 else 12)
        hpf.space_after = Pt(8)
        hrpr = hstyle.element.get_or_add_rPr()
        hea = hrpr.find(qn("w:rFonts"))
        if hea is None:
            hea = hstyle.element.makeelement(qn("w:rFonts"), {})
            hrpr.insert(0, hea)
        hea.set(qn("w:eastAsia"), KRDS_FONT)


def _add_footer(doc: Document, company_name: str = "") -> None:
    """Add KRDS-style footer with page number and company name."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if company_name:
            run = p.add_run(f"{company_name}  |  ")
            _set_font(run, 8, color=_GRAY_500)

        # Page number field
        run = p.add_run()
        _set_font(run, 8, color=_GRAY_500)
        fld_char1 = run._element.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "begin"})
        run._element.append(fld_char1)

        run2 = p.add_run()
        _set_font(run2, 8, color=_GRAY_500)
        instr = run2._element.makeelement(qn("w:instrText"), {})
        instr.text = " PAGE "
        run2._element.append(instr)

        run3 = p.add_run()
        _set_font(run3, 8, color=_GRAY_500)
        fld_char2 = run3._element.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "end"})
        run3._element.append(fld_char2)


def _add_cover_page(doc: Document, title: str, company_name: str = "") -> None:
    """Add KRDS-styled cover page with date."""
    # Spacer
    for _ in range(5):
        doc.add_paragraph()

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    _set_font(run, 28, bold=True, color=_BLUE_900)

    # Accent bar
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("━" * 20)
    _set_font(run2, 12, color=_BLUE_600)

    # Spacer
    for _ in range(2):
        doc.add_paragraph()

    # Company name
    if company_name:
        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run3 = p3.add_run(company_name)
        _set_font(run3, 16, bold=True, color=_GRAY_900)

    # Date (KST)
    kst = timezone(timedelta(hours=9))
    now = datetime.now(kst)
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_date = p_date.add_run(now.strftime("%Y년 %m월"))
    _set_font(run_date, 11, color=_GRAY_500)

    # Confidential
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run("CONFIDENTIAL")
    _set_font(run4, 9, color=_GRAY_500)

    doc.add_page_break()


def _add_toc_page(doc: Document, sections: list[tuple[str, str]]) -> None:
    """Add styled table of contents page."""
    doc.add_heading("목차", level=1)
    doc.add_paragraph()

    for i, (name, _) in enumerate(sections, 1):
        p = doc.add_paragraph()
        # Number in Blue 600
        run_num = p.add_run(f"{i:02d}  ")
        _set_font(run_num, 12, bold=True, color=_BLUE_600)
        # Section name
        run_name = p.add_run(name)
        _set_font(run_name, 12, color=_GRAY_900)
        # Dot leader + placeholder page
        run_dots = p.add_run(f"  {'·' * 30}  ")
        _set_font(run_dots, 10, color=_GRAY_200)

    doc.add_page_break()


# ---------------------------------------------------------------------------
# Markdown → DOCX conversion (with inline formatting preservation)
# ---------------------------------------------------------------------------

def _extract_text(children: list[dict[str, Any]]) -> str:
    """Recursively extract plain text from AST children (fallback only)."""
    parts: list[str] = []
    for child in children:
        if child["type"] == "text":
            parts.append(child.get("raw", child.get("children", "")))
        elif child["type"] == "codespan":
            parts.append(child.get("raw", child.get("children", "")))
        elif child["type"] in ("strong", "emphasis", "link"):
            parts.append(_extract_text(child.get("children", [])))
        elif isinstance(child.get("children"), list):
            parts.append(_extract_text(child["children"]))
        elif isinstance(child.get("raw"), str):
            parts.append(child["raw"])
    return "".join(parts)


def _add_inline_runs(paragraph, children: list[dict[str, Any]],
                     base_bold: bool = False, base_italic: bool = False) -> None:
    """Walk AST children and create runs with proper bold/italic formatting."""
    for child in children:
        ctype = child.get("type", "")

        if ctype == "text":
            text = child.get("raw", child.get("children", ""))
            if isinstance(text, str) and text:
                run = paragraph.add_run(text)
                run.bold = base_bold or None
                run.italic = base_italic or None

        elif ctype == "codespan":
            text = child.get("raw", child.get("children", ""))
            if isinstance(text, str) and text:
                run = paragraph.add_run(text)
                run.bold = base_bold or None
                run.italic = True
                run.font.color.rgb = _BLUE_800

        elif ctype == "strong":
            _add_inline_runs(paragraph, child.get("children", []),
                             base_bold=True, base_italic=base_italic)

        elif ctype == "emphasis":
            _add_inline_runs(paragraph, child.get("children", []),
                             base_bold=base_bold, base_italic=True)

        elif ctype == "link":
            _add_inline_runs(paragraph, child.get("children", []),
                             base_bold=base_bold, base_italic=base_italic)

        elif isinstance(child.get("children"), list):
            _add_inline_runs(paragraph, child["children"],
                             base_bold=base_bold, base_italic=base_italic)
        elif isinstance(child.get("raw"), str):
            text = child["raw"]
            if text:
                run = paragraph.add_run(text)
                run.bold = base_bold or None
                run.italic = base_italic or None


def _add_markdown_content(doc: Document, md_text: str) -> None:
    """Parse markdown via mistune AST and add to DOCX document."""
    tokens = _md_parser(md_text)
    if not isinstance(tokens, list):
        # fallback — if mistune returns raw HTML string, treat as plain text
        doc.add_paragraph(str(tokens))
        return

    for token in tokens:
        ttype = token.get("type", "")

        if ttype == "heading":
            level = min(token.get("attrs", {}).get("level", 1), 3)
            text = _extract_text(token.get("children", []))
            doc.add_heading(text, level=level)

        elif ttype == "paragraph":
            children = token.get("children", [])
            text = _extract_text(children)
            if text.strip():
                p = doc.add_paragraph()
                _add_inline_runs(p, children)

        elif ttype == "list":
            ordered = token.get("attrs", {}).get("ordered", False)
            style = "List Number" if ordered else "List Bullet"
            for item in token.get("children", []):
                item_children = item.get("children", [])
                # Collect all inline children from nested paragraphs
                all_inline: list[dict] = []
                for sub in item_children:
                    if sub.get("type") == "paragraph":
                        all_inline.extend(sub.get("children", []))
                    else:
                        all_inline.append(sub)
                text = _extract_text(all_inline)
                if text.strip():
                    p = doc.add_paragraph(style=style)
                    _add_inline_runs(p, all_inline)

        elif ttype == "block_code":
            code = token.get("raw", token.get("children", ""))
            if isinstance(code, str) and code.strip():
                doc.add_paragraph(code.strip())

        elif ttype == "thematic_break":
            doc.add_paragraph("─" * 40)

        else:
            # Fallback for unknown token types
            children = token.get("children", [])
            if isinstance(children, list):
                text = _extract_text(children)
            else:
                text = ""
            if not text:
                text = token.get("raw", "")
            if isinstance(text, str) and text.strip():
                doc.add_paragraph(text.strip())


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def assemble_docx(
    title: str,
    sections: list[tuple[str, str]],  # [(section_name, markdown_text), ...]
    output_path: str,
    author: str = "Kira Bot",
    company_name: str = "",
) -> str:
    """Assemble a DOCX proposal from section texts.

    Args:
        title: Proposal title.
        sections: List of (section_name, markdown_text) tuples.
        output_path: Where to save the .docx file.
        author: Document author metadata.
        company_name: Company name for cover page and footer.

    Returns:
        The output_path.
    """
    doc = Document()
    doc.core_properties.author = author

    # Apply KRDS styles
    _setup_styles(doc)
    _add_footer(doc, company_name)

    # Cover page
    _add_cover_page(doc, title, company_name)

    # Table of contents
    if sections:
        _add_toc_page(doc, sections)

    # Sections
    for name, content in sections:
        _add_markdown_content(doc, content)
        doc.add_page_break()

    doc.save(output_path)
    return output_path
