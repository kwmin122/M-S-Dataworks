"""Track Record Assembler — 실적/경력 기술서 DOCX 조립.

python-docx로 표(유사수행실적 요약표, 투입인력표) + 서술형 텍스트를 DOCX로 조립.

KRDS 디자인 시스템 기반 포맷:
- Pretendard 폰트 (없으면 맑은 고딕 fallback)
- Blue 900 (#003764) 기반 제목/헤딩
- Gray 700 (#444444) 본문
- 표지/목차/푸터 스타일링
"""
from __future__ import annotations

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from phase2_models import TrackRecordEntry, PersonnelEntry

# ---------------------------------------------------------------------------
# KRDS Design Tokens
# ---------------------------------------------------------------------------

_KRDS_FONT = "Pretendard"
_FALLBACK_FONT = "맑은 고딕"

_BLUE_900 = RGBColor(0x00, 0x37, 0x64)
_BLUE_800 = RGBColor(0x00, 0x4A, 0x8F)
_BLUE_600 = RGBColor(0x00, 0x70, 0xE0)
_BLUE_100 = RGBColor(0xE6, 0xF0, 0xFF)
_GRAY_900 = RGBColor(0x1A, 0x1A, 0x1A)
_GRAY_700 = RGBColor(0x44, 0x44, 0x44)
_GRAY_500 = RGBColor(0x88, 0x88, 0x88)
_GRAY_200 = RGBColor(0xE5, 0xE5, 0xE5)
_WHITE = RGBColor(0xFF, 0xFF, 0xFF)

_BLUE_900_HEX = "003764"
_BLUE_100_HEX = "E6F0FF"


def _set_font(run, size: float, bold: bool = False, color: RGBColor = _GRAY_700):
    """Apply KRDS font settings to a run."""
    run.font.name = _KRDS_FONT
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    ea = rpr.find(qn("w:rFonts"))
    if ea is None:
        ea = run._element.makeelement(qn("w:rFonts"), {})
        rpr.insert(0, ea)
    ea.set(qn("w:eastAsia"), _KRDS_FONT)


# ---------------------------------------------------------------------------
# Document Styles
# ---------------------------------------------------------------------------

def _setup_styles(doc: Document) -> None:
    """Configure document-level styles for KRDS compliance."""
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)   # 바인딩 여백
        section.right_margin = Cm(2.5)
        section.header_distance = Cm(1.2)
        section.footer_distance = Cm(1.0)

    style = doc.styles["Normal"]
    style.font.name = _KRDS_FONT
    style.font.size = Pt(11)
    style.font.color.rgb = _GRAY_700
    pf = style.paragraph_format
    pf.space_before = Pt(3)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.25
    rpr = style.element.get_or_add_rPr()
    ea = rpr.find(qn("w:rFonts"))
    if ea is None:
        ea = style.element.makeelement(qn("w:rFonts"), {})
        rpr.insert(0, ea)
    ea.set(qn("w:eastAsia"), _KRDS_FONT)

    for level in range(1, 4):
        hstyle = doc.styles[f"Heading {level}"]
        sizes = {1: 18, 2: 14, 3: 12}
        colors = {1: _BLUE_900, 2: _BLUE_800, 3: _GRAY_900}
        hstyle.font.name = _KRDS_FONT
        hstyle.font.size = Pt(sizes[level])
        hstyle.font.bold = True
        hstyle.font.color.rgb = colors[level]
        hpf = hstyle.paragraph_format
        hpf.space_before = Pt(18 if level == 1 else 12)
        hpf.space_after = Pt(8)
        hrpr = hstyle.element.get_or_add_rPr()
        hea = hrpr.find(qn("w:rFonts"))
        if hea is None:
            hea = hstyle.element.makeelement(qn("w:rFonts"), {})
            hrpr.insert(0, hea)
        hea.set(qn("w:eastAsia"), _KRDS_FONT)


def _add_footer(doc: Document, company_name: str = "") -> None:
    """Add KRDS-style footer with page number and company name."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if company_name:
            run = p.add_run(f"{company_name}  |  ")
            _set_font(run, 8, color=_GRAY_500)

        # Page number field
        run = p.add_run()
        _set_font(run, 8, color=_GRAY_500)
        fld_char1 = run._element.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "begin"})
        run._element.append(fld_char1)

        run2 = p.add_run()
        _set_font(run2, 8, color=_GRAY_500)
        instr = run2._element.makeelement(qn("w:instrText"), {})
        instr.text = " PAGE "
        run2._element.append(instr)

        run3 = p.add_run()
        _set_font(run3, 8, color=_GRAY_500)
        fld_char2 = run3._element.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "end"})
        run3._element.append(fld_char2)


def _add_cover_page(doc: Document, title: str, company_name: str = "") -> None:
    """Add KRDS-styled cover page with date."""
    from datetime import datetime, timezone, timedelta

    # Spacer
    for _ in range(5):
        doc.add_paragraph()

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    _set_font(run, 28, bold=True, color=_BLUE_900)

    # Accent bar
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("━" * 20)
    _set_font(run2, 12, color=_BLUE_600)

    # Spacer
    for _ in range(2):
        doc.add_paragraph()

    # Company name
    if company_name:
        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run3 = p3.add_run(company_name)
        _set_font(run3, 16, bold=True, color=_GRAY_900)

    # Date (KST)
    kst = timezone(timedelta(hours=9))
    now = datetime.now(kst)
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_date = p_date.add_run(now.strftime("%Y년 %m월"))
    _set_font(run_date, 11, color=_GRAY_500)

    # Confidential
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run("CONFIDENTIAL")
    _set_font(run4, 9, color=_GRAY_500)

    doc.add_page_break()


# ---------------------------------------------------------------------------
# Table Helpers
# ---------------------------------------------------------------------------

def _set_cell_text(cell, text: str, bold: bool = False, size: int = 9,
                   color: RGBColor = _GRAY_700) -> None:
    """테이블 셀에 KRDS 폰트로 텍스트 설정."""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    _set_font(run, size, bold=bold, color=color)


def _style_table_header(row, col_count: int) -> None:
    """KRDS Blue 900 배경 + 흰색 텍스트로 표 헤더 스타일링."""
    for i in range(col_count):
        cell = row.cells[i]
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{_BLUE_900_HEX}"/>')
        cell._tc.get_or_add_tcPr().append(shading)
        # Re-set text with white color
        existing_text = cell.text
        _set_cell_text(cell, existing_text, bold=True, size=9, color=_WHITE)


def _shade_row_alternate(row, row_index: int) -> None:
    """교대 행 배경색 (Blue 100) 설정."""
    if row_index % 2 == 0:
        for cell in row.cells:
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{_BLUE_100_HEX}"/>')
            cell._tc.get_or_add_tcPr().append(shading)


# ---------------------------------------------------------------------------
# Table Sections
# ---------------------------------------------------------------------------

def _add_track_record_summary_table(
    doc: Document,
    records: list[TrackRecordEntry],
) -> None:
    """유사수행실적 요약표."""
    doc.add_heading("유사수행실적 요약표", level=2)
    headers = ["No.", "프로젝트명", "발주처", "기간", "금액(억원)", "관련도"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        _set_cell_text(table.rows[0].cells[i], h, bold=True)
    _style_table_header(table.rows[0], len(headers))

    # Data rows
    for idx, rec in enumerate(records, 1):
        row = table.add_row()
        _set_cell_text(row.cells[0], str(idx))
        _set_cell_text(row.cells[1], rec.project_name[:30])
        _set_cell_text(row.cells[2], rec.client[:20])
        _set_cell_text(row.cells[3], rec.period or "-")
        amount_str = f"{rec.amount:.1f}" if rec.amount else "-"
        _set_cell_text(row.cells[4], amount_str)
        _set_cell_text(row.cells[5], f"{rec.relevance_score:.0%}")
        _shade_row_alternate(row, idx)

    doc.add_paragraph()  # spacing


def _add_track_record_details(
    doc: Document,
    records: list[TrackRecordEntry],
) -> None:
    """실적 상세 기술서."""
    doc.add_heading("유사수행실적 상세 기술서", level=2)
    for idx, rec in enumerate(records, 1):
        doc.add_heading(f"{idx}. {rec.project_name}", level=3)

        # Info line with Blue 600 accent
        info_items = []
        if rec.client:
            info_items.append(f"발주처: {rec.client}")
        if rec.period:
            info_items.append(f"기간: {rec.period}")
        if rec.amount:
            info_items.append(f"금액: {rec.amount:.1f}억원")
        if rec.technologies:
            info_items.append(f"기술: {', '.join(rec.technologies)}")
        if info_items:
            p = doc.add_paragraph()
            run = p.add_run(" | ".join(info_items))
            _set_font(run, 10, color=_BLUE_800)

        if rec.generated_text:
            doc.add_paragraph(rec.generated_text)
        elif rec.description:
            doc.add_paragraph(rec.description)

        doc.add_paragraph()  # spacing


def _add_personnel_table(
    doc: Document,
    personnel: list[PersonnelEntry],
) -> None:
    """투입인력 현황표."""
    doc.add_heading("투입인력 현황표", level=2)
    headers = ["No.", "성명", "역할", "등급", "경력(년)", "자격증"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        _set_cell_text(table.rows[0].cells[i], h, bold=True)
    _style_table_header(table.rows[0], len(headers))

    for idx, p in enumerate(personnel, 1):
        row = table.add_row()
        _set_cell_text(row.cells[0], str(idx))
        _set_cell_text(row.cells[1], p.name)
        _set_cell_text(row.cells[2], p.role)
        _set_cell_text(row.cells[3], p.grade or "-")
        _set_cell_text(row.cells[4], str(p.experience_years))
        certs = ", ".join(p.certifications[:3]) if p.certifications else "-"
        _set_cell_text(row.cells[5], certs)
        _shade_row_alternate(row, idx)

    doc.add_paragraph()


def _add_personnel_details(
    doc: Document,
    personnel: list[PersonnelEntry],
) -> None:
    """경력기술서 상세."""
    doc.add_heading("경력기술서", level=2)
    for idx, p in enumerate(personnel, 1):
        doc.add_heading(f"{idx}. {p.name} ({p.role})", level=3)

        info_items = []
        if p.grade:
            info_items.append(f"등급: {p.grade}")
        info_items.append(f"경력: {p.experience_years}년")
        if p.certifications:
            info_items.append(f"자격증: {', '.join(p.certifications)}")
        ip = doc.add_paragraph()
        run = ip.add_run(" | ".join(info_items))
        _set_font(run, 10, color=_BLUE_800)

        if p.generated_text:
            doc.add_paragraph(p.generated_text)

        if p.key_projects:
            doc.add_paragraph("주요 수행 프로젝트:")
            for proj in p.key_projects[:5]:
                doc.add_paragraph(f"  - {proj}", style="List Bullet")

        doc.add_paragraph()


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def assemble_track_record_docx(
    title: str,
    records: list[TrackRecordEntry],
    personnel: list[PersonnelEntry],
    output_path: str,
    author: str = "Kira Bot",
    company_name: str = "",
) -> str:
    """실적/경력 기술서 DOCX 조립.

    Args:
        title: 문서 제목.
        records: 유사수행실적 목록.
        personnel: 투입인력 목록.
        output_path: 출력 파일 경로.
        author: 문서 작성자 메타데이터.
        company_name: 회사명 (표지, 푸터에 표시).

    Returns:
        The output_path.
    """
    doc = Document()
    doc.core_properties.author = author

    # Apply KRDS styles
    _setup_styles(doc)
    _add_footer(doc, company_name)

    # Cover page
    _add_cover_page(doc, title, company_name)

    # Part 1: Track records
    if records:
        doc.add_heading("제1장 유사수행실적", level=1)
        _add_track_record_summary_table(doc, records)
        _add_track_record_details(doc, records)
        doc.add_page_break()

    # Part 2: Personnel
    if personnel:
        doc.add_heading("제2장 투입인력", level=1)
        _add_personnel_table(doc, personnel)
        _add_personnel_details(doc, personnel)

    doc.save(output_path)
    return output_path
